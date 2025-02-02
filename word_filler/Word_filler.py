import json
from docx import Document


def replace_placeholders_in_paragraph(paragraph, data):
    if not paragraph.runs:
        return
    else:
        for dicoo in data.values():#dico correspodant à "parties"
            for key, value in dicoo.items():#"bailleur" et dictionnaire corrspondzant
                if key in paragraph.text:
                    if "bool" not in value.keys() or value["bool"] == None:
                        paragraph.text = paragraph.text.replace(key, " ".join(str(v) for v in value.values() if v!=None))
                    elif value["bool"] ==True:
                        
                        paragraph.text = "Oui.  "+paragraph.text.replace(key, " ".join(str(v) for v in value.values() if v!=True and v!=None))
                    elif value["bool"] ==False:
                        paragraph.text = "Non.  "+paragraph.text.replace(key, " ".join(str(v) for v in value.values() if v!=False and v!=None))

                    
                    
def process_document(doc, data):
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    replace_placeholders_in_paragraph(paragraph, data)

def edit(donnees):
    with open(donnees, "r", encoding="utf-8") as json_file:
        data = json.load(json_file)
        data["AUTORISATIONS DE TRAVAUX"]["bailleur_travaux"] = data["AUTORISATIONS DE TRAVAUX"]["Bailleur"]["Travaux"]
        data["AUTORISATIONS DE TRAVAUX"]["bailleur_travaux_longue"] = data["AUTORISATIONS DE TRAVAUX"]["Bailleur"]["Travaux_longue_duree"]
        del data["AUTORISATIONS DE TRAVAUX"]["Bailleur"]
        data["AUTORISATIONS DE TRAVAUX"]["preneur_travaux"] = data["AUTORISATIONS DE TRAVAUX"]["Preneur"]["Travaux"]
        data["AUTORISATIONS DE TRAVAUX"]["preneur_plaques_enseignes"] = data["AUTORISATIONS DE TRAVAUX"]["Preneur"]["Plaques_enseignes"]
        del data["AUTORISATIONS DE TRAVAUX"]["Preneur"]
        

        data["HONORAIRES, IMPOTS, TAXES ET ASSURANCE DU BAILLEUR"]["parties_privees_bailleur"]=data["HONORAIRES, IMPOTS, TAXES ET ASSURANCE DU BAILLEUR"]["Parties_privees"]["Bailleur"]
        data["HONORAIRES, IMPOTS, TAXES ET ASSURANCE DU BAILLEUR"]["parties_privees_preneur"]=data["HONORAIRES, IMPOTS, TAXES ET ASSURANCE DU BAILLEUR"]["Parties_privees"]["Preneur"]
        del data["HONORAIRES, IMPOTS, TAXES ET ASSURANCE DU BAILLEUR"]["Parties_privees"]
        data["HONORAIRES, IMPOTS, TAXES ET ASSURANCE DU BAILLEUR"]["parties_communes_bailleur"]=data["HONORAIRES, IMPOTS, TAXES ET ASSURANCE DU BAILLEUR"]["Parties_communes"]["Bailleur"]
        data["HONORAIRES, IMPOTS, TAXES ET ASSURANCE DU BAILLEUR"]["parties_communes_preneur"]=data["HONORAIRES, IMPOTS, TAXES ET ASSURANCE DU BAILLEUR"]["Parties_communes"]["Preneur"]
        del data["HONORAIRES, IMPOTS, TAXES ET ASSURANCE DU BAILLEUR"]["Parties_communes"]
        
        
        data["TRAVAUX - REPARATIONS - REMPLACEMENTS"]["parties_privees_bailleur"]=data["TRAVAUX - REPARATIONS - REMPLACEMENTS"]["Parties_privees"]["Bailleur"]
        data["TRAVAUX - REPARATIONS - REMPLACEMENTS"]["parties_privees_preneur"]=data["TRAVAUX - REPARATIONS - REMPLACEMENTS"]["Parties_privees"]["Preneur"]
        del data["TRAVAUX - REPARATIONS - REMPLACEMENTS"]["Parties_privees"]
        data["TRAVAUX - REPARATIONS - REMPLACEMENTS"]["parties_communes_bailleur"]=data["TRAVAUX - REPARATIONS - REMPLACEMENTS"]["Parties_communes"]["Bailleur"]
        data["TRAVAUX - REPARATIONS - REMPLACEMENTS"]["parties_communes_preneur"]=data["TRAVAUX - REPARATIONS - REMPLACEMENTS"]["Parties_communes"]["Preneur"]
        del data["TRAVAUX - REPARATIONS - REMPLACEMENTS"]["Parties_communes"]
        
        data["HONORAIRES, IMPOTS, TAXES ET ASSURANCE DU BAILLEUR (2)"]["bailleur_gestion_tech"]=data["HONORAIRES, IMPOTS, TAXES ET ASSURANCE DU BAILLEUR (2)"]["Honoraires"]["Bailleur"]["Gestion_technique"]
        data["HONORAIRES, IMPOTS, TAXES ET ASSURANCE DU BAILLEUR (2)"]["bailleur_gestion_loc"]=data["HONORAIRES, IMPOTS, TAXES ET ASSURANCE DU BAILLEUR (2)"]["Honoraires"]["Bailleur"]["Gestion_locative"]
        data["HONORAIRES, IMPOTS, TAXES ET ASSURANCE DU BAILLEUR (2)"]["bailleur_gestion_loy"]=data["HONORAIRES, IMPOTS, TAXES ET ASSURANCE DU BAILLEUR (2)"]["Honoraires"]["Bailleur"]["Gestion_loyers"]
        data["HONORAIRES, IMPOTS, TAXES ET ASSURANCE DU BAILLEUR (2)"]["bailleur_gestion_syndic"]=data["HONORAIRES, IMPOTS, TAXES ET ASSURANCE DU BAILLEUR (2)"]["Honoraires"]["Bailleur"]["Gestion_syndic"]
        del data["HONORAIRES, IMPOTS, TAXES ET ASSURANCE DU BAILLEUR (2)"]["Honoraires"]["Bailleur"]

        data["HONORAIRES, IMPOTS, TAXES ET ASSURANCE DU BAILLEUR (2)"]["preneur_gestion_tech"]=data["HONORAIRES, IMPOTS, TAXES ET ASSURANCE DU BAILLEUR (2)"]["Honoraires"]["Preneur"]["Gestion_technique"]
        data["HONORAIRES, IMPOTS, TAXES ET ASSURANCE DU BAILLEUR (2)"]["preneur_gestion_loc"]=data["HONORAIRES, IMPOTS, TAXES ET ASSURANCE DU BAILLEUR (2)"]["Honoraires"]["Preneur"]["Gestion_locative"]
        data["HONORAIRES, IMPOTS, TAXES ET ASSURANCE DU BAILLEUR (2)"]["preneur_gestion_loy"]=data["HONORAIRES, IMPOTS, TAXES ET ASSURANCE DU BAILLEUR (2)"]["Honoraires"]["Preneur"]["Gestion_loyers"]
        data["HONORAIRES, IMPOTS, TAXES ET ASSURANCE DU BAILLEUR (2)"]["preneur_gestion_syndic"]=data["HONORAIRES, IMPOTS, TAXES ET ASSURANCE DU BAILLEUR (2)"]["Honoraires"]["Preneur"]["Gestion_syndic"]
        del data["HONORAIRES, IMPOTS, TAXES ET ASSURANCE DU BAILLEUR (2)"]["Honoraires"]["Preneur"]
        
        return data
        


    

def fill(path_of_data):
    data=edit(path_of_data)
    doc = Document("Fiche_audit_modele.docx")
    
    # Ajout des placeholders avec accolades
    for dictionnaire in data.values():
        for key in list(dictionnaire.keys()):  # Utilisation de list() pour éviter des problèmes lors de la modification
            dictionnaire[f"{{{{{key}}}}}"] = dictionnaire.pop(key)
    
    process_document(doc, data)
    doc.save("fiche_audit_remplie.docx")

if __name__ == "__main__":
    fill("data.json")
